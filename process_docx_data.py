#!/usr/bin/env python3
"""
Process DOCX files to extract work hours and payment data for the analytics dashboard.
"""

import json
import os
from docx import Document
import re

def extract_data_from_docx(file_path):
    """Extract hours, earnings, and payment data from a DOCX file."""
    try:
        doc = Document(file_path)
        text = ""
        
        # Extract all text from the document
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        print(f"Extracted text from {file_path}:")
        print(text)
        print("-" * 50)
        
        # Initialize data
        hours = 0
        amount = 0
        received = 0
        
        # Patterns to extract data
        hours_patterns = [
            r'total.*?hours.*?(\d+\.?\d*)',
            r'hours.*?worked.*?(\d+\.?\d*)',
            r'(\d+\.?\d*).*?hours',
            r'hours.*?(\d+\.?\d*)'
        ]
        
        amount_patterns = [
            r'total.*?amount.*?\$?(\d+,?\d*\.?\d*)',
            r'earned.*?\$?(\d+,?\d*\.?\d*)',
            r'amount.*?earned.*?\$?(\d+,?\d*\.?\d*)',
            r'\$(\d+,?\d*\.?\d*)'
        ]
        
        received_patterns = [
            r'received.*?\$?(\d+,?\d*\.?\d*)',
            r'payment.*?\$?(\d+,?\d*\.?\d*)',
            r'paid.*?\$?(\d+,?\d*\.?\d*)'
        ]
        
        # Try to extract hours
        for pattern in hours_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    hours = float(match.group(1))
                    break
                except ValueError:
                    continue
        
        # Try to extract amount
        for pattern in amount_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    # Remove commas and convert to float
                    clean_amount = float(match.replace(',', ''))
                    if clean_amount > amount:  # Take the largest amount found
                        amount = clean_amount
                except ValueError:
                    continue
        
        # Try to extract received amount
        for pattern in received_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    # Remove commas and convert to float
                    clean_received = float(match.replace(',', ''))
                    if clean_received > received:  # Take the largest received amount found
                        received = clean_received
                except ValueError:
                    continue
        
        return {
            'hours': hours,
            'amount': amount,
            'received': received,
            'raw_text': text[:500]  # Store first 500 chars for debugging
        }
        
    except Exception as e:
        print(f"Error processing {file_path}: {str(e)}")
        return None

def update_dashboard_data():
    """Update the dashboard with new data from DOCX files."""
    
    # Current data from the dashboard
    current_data = [
        {'month': 'March 2024', 'hours': 142.25, 'amount': 3556.25, 'received': 0, 'notes': 'Amount to be received: $14,698.67'},
        {'month': 'April 2024', 'hours': 192.58, 'amount': 4814.5, 'received': 1000, 'notes': ''},
        {'month': 'May 2024', 'hours': 213.67, 'amount': 5341.67, 'received': 1200, 'notes': ''},
        {'month': 'June 2024', 'hours': 224.50, 'amount': 5612.5, 'received': 6170, 'notes': ''},
        {'month': 'July 2024', 'hours': 176.0, 'amount': 4400.0, 'received': 2000, 'notes': 'Previous Amount to be received before April= 13,098.67'},
        {'month': 'August 2024', 'hours': 231.0, 'amount': 5775.0, 'received': 3500, 'notes': 'Amount received in April= 5,000 + 11,000= 16,000'},
        {'month': 'September 2024', 'hours': 202.75, 'amount': 5068.75, 'received': 2700, 'notes': 'Total hours April 1-15= 140'},
        {'month': 'October 2024', 'hours': 178.25, 'amount': 4456.25, 'received': 14775, 'notes': 'Total amount April 1-15= 140x25= 3,500'},
        {'month': 'November 2024', 'hours': 176.25, 'amount': 4406.25, 'received': 2000, 'notes': ''},
        {'month': 'December 2024', 'hours': 222.0, 'amount': 5550.0, 'received': 500, 'notes': 'Total amount to be received till April= (13,098.67 + 3,500)- 16,000 = 598.67'},
        {'month': 'January 2025', 'hours': 249.5, 'amount': 6237.5, 'received': 1400, 'notes': ''},
        {'month': 'February 2025', 'hours': 222.75, 'amount': 5568.75, 'received': 16750, 'notes': ''},
        {'month': 'March 2025', 'hours': 236.25, 'amount': 5906.25, 'received': 1600, 'notes': ''},
        {'month': 'April 2025', 'hours': 140.0, 'amount': 3500.0, 'received': 16000, 'notes': ''},
        {'month': 'May 2025', 'hours': 0.0, 'amount': 0.0, 'received': 6000, 'notes': ''},
        {'month': 'June 2025', 'hours': 226.0, 'amount': 6102.0, 'received': 4000, 'notes': ''},
        {'month': 'July 2025', 'hours': 206.75, 'amount': 5582.25, 'received': 3150, 'notes': ''},
        {'month': 'August 2025', 'hours': 202.75, 'amount': 5474.25, 'received': 4240, 'notes': ''}
    ]
    
    # Process new DOCX files
    docx_files = {
        'September 2025': '2025 Sept 1-30.docx',
        'October 2025': '2025 Oct 1-31.docx'
    }
    
    new_entries = []
    
    for month, filename in docx_files.items():
        if os.path.exists(filename):
            print(f"Processing {filename} for {month}...")
            data = extract_data_from_docx(filename)
            if data:
                new_entry = {
                    'month': month,
                    'hours': data['hours'],
                    'amount': data['amount'],
                    'received': data['received'],
                    'notes': f'Data extracted from {filename}'
                }
                new_entries.append(new_entry)
                print(f"Extracted data for {month}: {data}")
            else:
                print(f"Failed to extract data from {filename}")
        else:
            print(f"File {filename} not found")
    
    # Update current data with new entries
    # Remove existing September and October 2025 entries if they exist
    current_data = [item for item in current_data if item['month'] not in ['September 2025', 'October 2025']]
    
    # Add new entries
    current_data.extend(new_entries)
    
    # Sort by date
    def sort_key(item):
        month_name, year = item['month'].split(' ')
        months = ['January', 'February', 'March', 'April', 'May', 'June',
                 'July', 'August', 'September', 'October', 'November', 'December']
        try:
            month_index = months.index(month_name)
        except ValueError:
            month_index = 0
        return (int(year), month_index)
    
    current_data.sort(key=sort_key)
    
    # Save updated data to JSON file
    with open('data.json', 'w') as f:
        json.dump(current_data, f, indent=2)
    
    print(f"\nUpdated data.json with {len(new_entries)} new entries")
    print("Summary of all data:")
    for item in current_data:
        print(f"{item['month']}: {item['hours']} hours, ${item['amount']}, received ${item['received']}")
    
    return current_data

if __name__ == "__main__":
    updated_data = update_dashboard_data()